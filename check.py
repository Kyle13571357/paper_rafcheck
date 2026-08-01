#!/usr/bin/env python3
"""Audit written prose against the papers it cites.

Four steps, in order:

  1. Claim extraction   - pull out {subject, metric, value, unit, relation,
                          condition, cited_ref, source_span}. The `condition`
                          field is not optional: without it the most valuable
                          finding (a correct number whose qualifying condition
                          was dropped) is undetectable by construction.
  2. Reference resolution - "[9]" -> m5_asplos25, via corpus.yaml.
  3. Scope decision     - per claim, never per paragraph. One sentence citing
                          three different papers runs three scoped searches.
  4. Verdict + span     - every result carries the retrieved original text, so
                          a human can check the tool rather than trust it.

The scope rule that makes this work at all: when a claim has no citation and
looks like a restatement of someone else's result, the search runs globally
but EXCLUDES tier-0. Allowing tier-0 would let the survey corroborate itself
-- the system would appear to function while verifying paraphrase against
paraphrase.

Usage:
  python3 check.py --text "Colloid achieves 1.01-1.76x speedup [3]."
  python3 check.py --file draft.docx        # whole draft, .docx/.pdf/.txt
  python3 check.py --file draft.docx --all  # include claims that checked out
  python3 check.py --self-audit             # audit the survey against tier-1
  python3 check.py --self-audit --dry-run   # scope decisions only, no model
"""

import argparse
import json
import re
import sys
from pathlib import Path

import yaml

import llm
from units import quantities_match, parse_quantity

from corpus import load_corpus

ROOT = Path(__file__).parent

VERDICTS = ("supported", "contradicted", "not_found",
            "condition_mismatch", "underspecified")

EXTRACT_PROMPT = """Extract every verifiable factual claim from the passage below.

A verifiable claim states something checkable about a system: a measurement, a
comparison, a capability, a definition. Ignore opinions, motivation and
transitions.

For each claim record:
  subject      - the system or thing the claim is about (e.g. "Colloid")
  metric       - what is measured (e.g. "performance improvement")
  value        - the number or value exactly as written ("1.01-1.76x"), or null
  unit         - the unit if separable, else null
  relation     - how value relates to subject: "equals", "up_to", "at_least",
                 "improves_over", "range", "defines", or "qualitative"
  condition    - EVERY qualifying condition attached to the claim: workload,
                 configuration, baseline, hardware, scale. Null ONLY if the
                 passage truly states none. This field matters more than the
                 others: a number quoted without its condition is the most
                 common way a survey distorts a source.
  cited_ref    - the bracketed citation number attached to this claim, as an
                 integer (from "[9]" record 9). Null if the claim carries none.
  source_span  - the exact substring of the passage this claim comes from,
                 copied verbatim so it can be located again.
  is_own_result - true if the passage presents this as the authors' OWN
                 experiment or system, false if it restates someone else's work.

Reply with ONLY a JSON object: {"claims": [ ... ]}"""

ADJUDICATE_PROMPT = """You are checking one claim against passages retrieved from the source it cites.

Choose exactly one verdict:

  supported          - the sources state this value AND the same conditions
  contradicted       - the sources state a clearly different value for the same thing
  not_found          - the sources do not state this at all (wrong paper cited, or
                       the writer derived the value themselves)
  condition_mismatch - the VALUE matches but a qualifying condition was dropped,
                       weakened or changed. Use this whenever the number is right
                       but the sources attach a workload, baseline, scale or
                       configuration that the claim omits.
  underspecified     - the sources themselves are ambiguous or inconsistent: the
                       same symbol defined twice, a condition stated two ways,
                       a relationship left unexplained. Use this rather than
                       picking a side.

Rules:
  - Judge only from the passages provided. Absence of evidence is not_found,
    never contradicted.
  - Do not do arithmetic to make a claim fit. A value the writer computed from
    source numbers is not_found, however reasonable the computation.
  - Quote the exact sentence you relied on in evidence_span. If you cannot
    quote one, the verdict cannot be "supported".

Reply with ONLY a JSON object:
{
  "verdict": "<one of the five>",
  "confidence": "high" | "medium" | "low",
  "evidence_span": "<verbatim sentence from the passages, or empty>",
  "explanation": "<one or two sentences>",
  "condition_in_source": "<the condition the sources attach, or empty>"
}"""


def build_ref_map(corpus):
    """ref number -> entry. This table, not the model, resolves citations."""
    return corpus.by_ref


def resolve_reference(cited_ref, ref_map):
    """Returns (doc_ids, status). Handles compound citations like [3, 9]."""
    if cited_ref is None:
        return [], "no_citation"
    refs = cited_ref if isinstance(cited_ref, list) else [cited_ref]
    doc_ids, unknown = [], []
    for r in refs:
        try:
            r = int(r)
        except (TypeError, ValueError):
            unknown.append(r)
            continue
        entry = ref_map.get(r)
        if entry:
            doc_ids.append(entry["doc_id"])
        else:
            unknown.append(r)
    if unknown and not doc_ids:
        return [], f"unresolvable_ref:{unknown}"
    if unknown:
        return doc_ids, f"partially_resolved:{unknown}"
    return doc_ids, "resolved"


def decide_scope(claim, ref_map):
    """The scope decision from the proposal, made per claim.

    Returns a dict describing what to search and why, so the reasoning is
    inspectable in the output rather than buried in control flow."""
    doc_ids, status = resolve_reference(claim.get("cited_ref"), ref_map)

    if doc_ids:
        return {
            "action": "verify",
            "doc_ids": doc_ids,
            "exclude_tier": None,
            "reason": f"claim cites {claim.get('cited_ref')} -> {doc_ids}; "
                      f"search locked to that work",
            "ref_status": status,
        }

    if status.startswith("unresolvable_ref"):
        return {
            "action": "flag",
            "doc_ids": [], "exclude_tier": None,
            "reason": f"citation {claim.get('cited_ref')} is not in corpus.yaml",
            "ref_status": status,
        }

    if claim.get("is_own_result"):
        return {
            "action": "skip",
            "doc_ids": [], "exclude_tier": None,
            "reason": "presented as the authors' own result; nothing to verify against",
            "ref_status": status,
        }

    return {
        "action": "verify_uncited",
        "doc_ids": [],
        "exclude_tier": 0,
        "reason": "uncited restatement of others' work; global search with tier-0 "
                  "excluded so the document cannot corroborate itself",
        "ref_status": status,
    }


def extract_claims(passage, timeout=300):
    raw = llm.complete(f"{EXTRACT_PROMPT}\n\n--- PASSAGE ---\n{passage}\n",
                       json_mode=True, temperature=0.0, timeout=timeout)
    data = llm.parse_json_reply(raw) or {}
    claims = data.get("claims") or []
    out = []
    for c in claims:
        if not isinstance(c, dict):
            continue
        span = (c.get("source_span") or "").strip()
        out.append({
            "subject": c.get("subject"),
            "metric": c.get("metric"),
            "value": c.get("value"),
            "unit": c.get("unit"),
            "relation": c.get("relation"),
            "condition": c.get("condition"),
            "cited_ref": c.get("cited_ref"),
            "source_span": span,
            "span_verified": bool(span) and span in passage,
            "is_own_result": bool(c.get("is_own_result")),
        })
    return out


def format_evidence(hits):
    parts = []
    for i, h in enumerate(hits, 1):
        pages = (f"p{h['page_start']}" if h["page_start"] == h["page_end"]
                 else f"p{h['page_start']}-{h['page_end']}")
        parts.append(f"[E{i}] {h['doc_id']} {pages} "
                     f"(section: {h.get('section') or 'n/a'})\n{h['text']}")
    return "\n\n".join(parts)


# A number whose unit was destroyed by a broken font table. parse.py maps the
# Mathematical Alphanumeric range to ASCII, so "54 µs" ends up as "54ZM" -- it
# reads as ordinary text and matches no query naming the unit.
MISENCODED_UNIT_RE = re.compile(r"\d\s*(?:[\U0001D400-\U0001D7FF]{1,4}|[A-Z]{2,4}\b)")


def numeric_prescreen(claim, hits):
    """Deterministic check of whether the claimed number appears at all.

    Runs before the model so the model's verdict can be compared against
    arithmetic that did not come from a model. Returns None when the claim has
    no comparable number.

    A negative result is reported as *unreliable* when the retrieved text
    contains numbers whose units were lost to a broken font encoding. That
    distinction matters: searching M5 for "54 µs" finds nothing, yet the paper
    states it plainly -- the unit extracts as "ZM". Reporting a confident
    "not found" there produces a false accusation against a correct citation,
    which is precisely the error this project exists to prevent."""
    claimed = claim.get("value")
    if claimed is None or parse_quantity(claimed) is None:
        return None
    suspect = []
    for h in hits:
        for sent in re.split(r"(?<=[.!?])\s+", h["text"]):
            for m in re.finditer(r"[-+]?\d[\d,]*(?:\.\d+)?\s*"
                                 r"(?:%|×|x|[a-zA-Zµμ]+(?:/[a-zA-Z]+)?)?", sent):
                if quantities_match(claimed, m.group(0)):
                    return {"found": True, "matched": m.group(0).strip(),
                            "in_doc": h["doc_id"],
                            "page": h["page_start"], "sentence": sent.strip()[:300]}
        for m in MISENCODED_UNIT_RE.finditer(h["text"]):
            suspect.append({"doc_id": h["doc_id"], "page": h["page_start"],
                            "text": m.group(0)})

    result = {"found": False}
    if suspect:
        result["reliable"] = False
        result["reason"] = (
            f"{len(suspect)} number(s) in the retrieved text carry units lost to a "
            f"broken font encoding (e.g. {suspect[0]['text']!r} in "
            f"{suspect[0]['doc_id']} p{suspect[0]['page']}), so a unit-bearing "
            f"value can be present yet unmatchable. Check the rendered page "
            f"before concluding the figure is absent.")
        result["sites"] = suspect[:5]
    return result


def adjudicate(claim, hits, timeout=300):
    claim_desc = json.dumps({k: claim[k] for k in
                             ("subject", "metric", "value", "unit",
                              "relation", "condition")}, ensure_ascii=False)
    prompt = (f"{ADJUDICATE_PROMPT}\n\n--- CLAIM ---\n{claim_desc}\n"
              f"\nClaim as written: {claim['source_span']}\n"
              f"\n--- RETRIEVED PASSAGES ---\n{format_evidence(hits)}\n")
    raw = llm.complete(prompt, json_mode=True, temperature=0.0, timeout=timeout)
    data = llm.parse_json_reply(raw) or {}
    verdict = str(data.get("verdict", "")).strip().lower()
    if verdict not in VERDICTS:
        verdict = "underspecified"
    return {
        "verdict": verdict,
        "confidence": str(data.get("confidence", "")).strip().lower() or "low",
        "evidence_span": (data.get("evidence_span") or "").strip(),
        "explanation": (data.get("explanation") or "").strip(),
        "condition_in_source": (data.get("condition_in_source") or "").strip(),
    }


def check_passage(passage, retriever=None, k=6, dry_run=False, claims=None,
                  corpus=None):
    corpus = corpus or (retriever.corpus if retriever is not None else load_corpus())
    ref_map = build_ref_map(corpus)

    if claims is None:
        if dry_run:
            raise ValueError("--dry-run needs claims supplied via --claims-json")
        claims = extract_claims(passage)

    from retrieve import Retriever
    r = retriever
    if r is None and not dry_run:
        r = Retriever()

    results = []
    for claim in claims:
        scope = decide_scope(claim, ref_map)
        record = {"claim": claim, "scope": scope}

        if scope["action"] in ("skip", "flag"):
            record["verdict"] = ("skipped_own_result" if scope["action"] == "skip"
                                 else "unresolvable_reference")
            record["evidence"] = []
            results.append(record)
            continue

        if dry_run:
            record["verdict"] = "(dry-run: not adjudicated)"
            record["evidence"] = []
            results.append(record)
            continue

        query = " ".join(str(x) for x in
                         (claim.get("subject"), claim.get("metric"),
                          claim.get("value"), claim.get("condition")) if x)
        hits = r.search(query, k=k,
                        doc_id=scope["doc_ids"] or None,
                        exclude_tier=scope["exclude_tier"])

        record["evidence"] = [{
            "doc_id": h["doc_id"], "tier": h["tier"],
            "page_start": h["page_start"], "page_end": h["page_end"],
            "section": h.get("section"), "chunk_id": h["chunk_id"],
            "text": h["text"],
        } for h in hits]
        record["numeric_prescreen"] = numeric_prescreen(claim, hits)

        if not hits:
            record.update({"verdict": "not_found", "confidence": "high",
                           "explanation": "retrieval returned nothing in scope"})
            results.append(record)
            continue

        verdict = adjudicate(claim, hits)
        record.update(verdict)

        # A "not_found" resting on text whose units were destroyed by a broken
        # font is not a finding, it is a limitation of the extraction. Downgrade
        # it to needs-human rather than let it stand as an accusation against a
        # citation that may well be correct.
        ps = record.get("numeric_prescreen") or {}
        if (record["verdict"] == "not_found"
                and not ps.get("found") and ps.get("reliable") is False):
            record["verdict"] = "underspecified"
            record["confidence"] = "low"
            record["explanation"] = (
                "cannot verify from the text layer: " + ps["reason"] + " "
                + (verdict.get("explanation") or ""))
            record["needs_visual_check"] = True

        # An uncited claim whose number *is* present in some original is the
        # "possibly missing citation" case, not a plain pass.
        if (scope["action"] == "verify_uncited"
                and verdict["verdict"] == "supported"
                and record["numeric_prescreen"]
                and record["numeric_prescreen"].get("found")):
            record["verdict"] = "possibly_missing_citation"
            record["explanation"] = (
                f"value found in {record['numeric_prescreen']['in_doc']} "
                f"p{record['numeric_prescreen']['page']} but the passage cites no "
                f"reference. " + verdict["explanation"])

        results.append(record)

    return {"passage": passage, "model": llm.describe() if not dry_run else None,
            "results": results}


def render(report, verbose=False):
    print(f"claims examined: {len(report['results'])}")
    if report.get("model"):
        print(f"model: {report['model']}")
    print()
    counts = {}
    for rec in report["results"]:
        v = rec.get("verdict", "?")
        counts[v] = counts.get(v, 0) + 1
        claim = rec["claim"]
        print(f"[{v}] {claim.get('subject')} - {claim.get('metric')}"
              f" = {claim.get('value')}")
        print(f"    span    : {claim.get('source_span')!r}")
        if claim.get("condition"):
            print(f"    condition: {claim['condition']}")
        print(f"    scope   : {rec['scope']['reason']}")
        if rec.get("condition_in_source"):
            print(f"    source condition: {rec['condition_in_source']}")
        if rec.get("explanation"):
            print(f"    why     : {rec['explanation']}")
        ps = rec.get("numeric_prescreen")
        if ps is not None:
            if ps.get("found"):
                print(f"    numeric : {ps['matched']!r} found in {ps['in_doc']} "
                      f"p{ps['page']} (deterministic)")
            else:
                print("    numeric : claimed value not found by exact match "
                      "(deterministic)")
        if rec.get("evidence_span"):
            print(f"    evidence: {rec['evidence_span'][:220]}")
        elif verbose and rec.get("evidence"):
            e = rec["evidence"][0]
            print(f"    top hit : {e['doc_id']} p{e['page_start']}: "
                  f"{e['text'][:200]}")
        print()
    print("summary:", ", ".join(f"{k}={v}" for k, v in sorted(counts.items())))


# ---------------------------------------------------------------------------
# Whole-document proofreading
# ---------------------------------------------------------------------------

#: worst first -- a proofreading pass should open on what is actually wrong
VERDICT_PRIORITY = {
    "contradicted": 0,
    "not_found": 1,
    "condition_mismatch": 2,
    "underspecified": 3,
    "possibly_missing_citation": 4,
    "unresolvable_reference": 5,
    "supported": 6,
    "skipped_own_result": 7,
}
PROBLEM_VERDICTS = {
    "contradicted", "not_found", "condition_mismatch",
    "underspecified", "possibly_missing_citation", "unresolvable_reference",
}


def read_document(path):
    """Draft text from .docx, .pdf or plain text.

    .docx first because that is what the survey is actually written in;
    proofreading should not require exporting to text by hand.
    """
    p = Path(path)
    suffix = p.suffix.lower()
    if suffix == ".docx":
        import docx
        doc = docx.Document(str(p))
        parts = [para.text for para in doc.paragraphs]
        for table in doc.tables:            # numbers in tables are claims too
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    if suffix == ".pdf":
        import fitz
        d = fitz.open(str(p))
        text = "\n".join(page.get_text() for page in d)
        d.close()
        return text
    return p.read_text()


def _bounded_pieces(block, max_chars):
    """Cut an oversized block into pieces, preferring the most meaningful
    boundary available.

    Order matters: sentences first, then table rows (a flattened table has no
    sentence punctuation at all, so it would otherwise stay one 4000-character
    lump), then a hard wrap so the function always terminates."""
    if len(block) <= max_chars:
        return [block]

    for splitter in (r"(?<=[.!?])\s+(?=[A-Z(\[])", r"\n"):
        units = [u for u in re.split(splitter, block) if u.strip()]
        if len(units) < 2:
            continue
        pieces, cur = [], ""
        for u in units:
            if cur and len(cur) + len(u) + 1 > max_chars:
                pieces.append(cur)
                cur = u
            else:
                cur = f"{cur} {u}".strip() if splitter.startswith("(?<=") else \
                      f"{cur}\n{u}".strip()
            # a single unit longer than the cap still needs cutting
            while len(cur) > max_chars:
                cut = cur.rfind(" ", 0, max_chars)
                cut = cut if cut > max_chars // 2 else max_chars
                pieces.append(cur[:cut])
                cur = cur[cut:].strip()
        if cur:
            pieces.append(cur)
        if all(len(p) <= max_chars for p in pieces):
            return pieces

    return [block[i:i + max_chars] for i in range(0, len(block), max_chars)]


def split_paragraphs(text, min_chars=80, max_chars=1500):
    """Paragraphs worth checking.

    Claim extraction needs enough context to see the condition attached to a
    number, so very short lines (headings, captions, list bullets) are merged
    forward rather than sent on their own. Newlines are preserved until after
    splitting so table rows remain usable cut points."""
    raw = [b.strip() for b in re.split(r"\n\s*\n|\r\n\r\n", text) if b.strip()]

    merged, buf = [], ""
    for b in raw:
        if len(buf) + len(b) + 1 <= max_chars and len(buf) < min_chars:
            buf = f"{buf}\n{b}".strip()
        else:
            if buf:
                merged.append(buf)
            buf = b
    if buf:
        merged.append(buf)

    out = []
    for m in merged:
        for piece in _bounded_pieces(m, max_chars):
            piece = re.sub(r"\s+", " ", piece).strip()
            # only text containing a digit can carry a numeric claim; the rest
            # would cost a model call to conclude nothing
            if len(piece) >= min_chars and re.search(r"\d", piece):
                out.append(piece)
    return out


def check_document(path, retriever=None, k=6, limit=None, progress=True):
    # One clear message beats the same credentials error repeated once per
    # paragraph after the user has already waited for retrieval to load.
    llm.require_key(llm.config())
    text = read_document(path)
    paras = split_paragraphs(text)
    if limit:
        paras = paras[:limit]
    if progress:
        print(f"{Path(path).name}: {len(paras)} 個含數字的段落待檢查\n",
              file=sys.stderr)

    all_results = []
    for i, para in enumerate(paras, 1):
        if progress:
            print(f"  [{i}/{len(paras)}] …", file=sys.stderr)
        try:
            rep = check_passage(para, retriever=retriever, k=k)
        except Exception as e:                          # noqa: BLE001
            print(f"  段落 {i} 失敗: {type(e).__name__}: {e}", file=sys.stderr)
            continue
        for rec in rep["results"]:
            rec["paragraph_index"] = i
            rec["paragraph"] = para
            all_results.append(rec)
    return {"source": str(path), "paragraphs": len(paras),
            "model": llm.describe(), "results": all_results}


def render_document_report(report, show_all=False):
    results = report["results"]
    problems = [r for r in results if r.get("verdict") in PROBLEM_VERDICTS]
    shown = results if show_all else problems
    shown = sorted(shown, key=lambda r: (VERDICT_PRIORITY.get(r.get("verdict"), 9),
                                         r.get("paragraph_index", 0)))

    counts = {}
    for r in results:
        v = r.get("verdict", "?")
        counts[v] = counts.get(v, 0) + 1

    print("=" * 78)
    print(f"校對報告: {report['source']}")
    print(f"段落 {report['paragraphs']} · claim {len(results)} · "
          f"需注意 {len(problems)}")
    print(f"model: {report['model']}")
    print("=" * 78)

    if not shown:
        print("\n沒有發現問題。")
    for r in shown:
        c = r["claim"]
        print(f"\n【{r.get('verdict')}】 段落 {r.get('paragraph_index')}"
              f"  ·  {c.get('subject')} — {c.get('metric')} = {c.get('value')}")
        print(f"  你寫的   : {c.get('source_span')}")
        if c.get("condition"):
            print(f"  你的條件 : {c['condition']}")
        if r.get("condition_in_source"):
            print(f"  原文條件 : {r['condition_in_source']}")
        if r.get("explanation"):
            print(f"  說明     : {r['explanation']}")
        ps = r.get("numeric_prescreen")
        if ps is not None and not ps.get("found"):
            print("  數值比對 : 引用來源中找不到此數值（程式判定,非模型）")
        elif ps and ps.get("found"):
            print(f"  數值比對 : {ps['matched']!r} 見 {ps['in_doc']} "
                  f"p{ps['page']}（程式判定）")
        if r.get("evidence_span"):
            print(f"  原文     : {r['evidence_span'][:300]}")
        elif r.get("evidence"):
            e = r["evidence"][0]
            print(f"  最相近   : {e['doc_id']} p{e['page_start']} — "
                  f"{e['text'][:220]}")

    print("\n" + "-" * 78)
    print("統計: " + ", ".join(f"{k}={v}" for k, v in
                               sorted(counts.items(),
                                      key=lambda kv: VERDICT_PRIORITY.get(kv[0], 9))))
    if not show_all and len(results) > len(problems):
        print(f"（已隱藏 {len(results) - len(problems)} 筆沒問題的 claim,"
              f"用 --all 顯示）")


def self_audit_passages(corpus, limit=8):
    """Passages from the tier-0 document to audit against what they cite.

    Prefers the list declared in checks.yaml. With none declared, samples
    cited passages out of the tier-0 document itself, so `--self-audit` works
    on a corpus nobody has hand-curated."""
    declared = corpus.checks.get("self_audit") or []
    if declared:
        return [re.sub(r"\s+", " ", p).strip() for p in declared]

    tier0 = corpus.primary_tier0()
    if tier0 is None:
        return []
    import json as _json
    passages = []
    with open(corpus.blocks_path) as f:
        for line in f:
            b = _json.loads(line)
            if b["doc_id"] != tier0["doc_id"] or b["block_type"] != "prose":
                continue
            # a passage is worth auditing when it carries both a citation and
            # a number -- that is what check.py can actually verify
            if re.search(r"\[\d+\]", b["text"]) and re.search(r"\d", b["text"]):
                passages.append(re.sub(r"\s+", " ", b["text"]).strip())
    passages.sort(key=len, reverse=True)
    return passages[:limit]


def main():
    ap = argparse.ArgumentParser()
    g = ap.add_mutually_exclusive_group(required=True)
    g.add_argument("--text")
    g.add_argument("--file")
    g.add_argument("--self-audit", action="store_true")
    ap.add_argument("-k", type=int, default=6)
    ap.add_argument("--dry-run", action="store_true",
                    help="show scope decisions without calling any model")
    ap.add_argument("--claims-json", help="pre-extracted claims (for --dry-run)")
    ap.add_argument("--json", action="store_true")
    ap.add_argument("--verbose", action="store_true")
    ap.add_argument("--out", help="write the full report to this path")
    ap.add_argument("--all", action="store_true",
                    help="document mode: also show claims that checked out")
    ap.add_argument("--limit", type=int, default=None,
                    help="document mode: only the first N paragraphs")
    ap.add_argument("--corpus", default=None, help="corpus directory")
    args = ap.parse_args()

    # a whole draft goes through the document path: paragraph splitting plus a
    # report ordered by severity, rather than one wall of output
    if args.file and not args.dry_run:
        try:
            # before the 16s retrieval load, not after
            llm.require_key(llm.config())
        except llm.LLMError as e:
            print(f"error: {e}", file=sys.stderr)
            return 1
        from retrieve import Retriever
        report = check_document(args.file, retriever=Retriever(load_corpus(args.corpus)),
                                k=args.k, limit=args.limit)
        if args.json:
            print(json.dumps(report, indent=2, ensure_ascii=False))
        else:
            render_document_report(report, show_all=args.all)
        if args.out:
            Path(args.out).write_text(json.dumps(report, indent=2,
                                                 ensure_ascii=False))
            print(f"\nwrote {args.out}", file=sys.stderr)
        return

    corpus = load_corpus(args.corpus)

    if args.self_audit:
        passages = self_audit_passages(corpus)
        if not passages:
            print("no tier-0 document registered, and checks.yaml declares no "
                  "self_audit passages -- nothing to audit", file=sys.stderr)
            return 1
    elif args.file:
        passages = [read_document(args.file)]
    else:
        passages = [args.text]

    preset = None
    if args.claims_json:
        preset = json.loads(Path(args.claims_json).read_text())

    retriever = None
    if not args.dry_run:
        from retrieve import Retriever
        retriever = Retriever(corpus)

    reports = []
    for i, p in enumerate(passages):
        claims = None
        if preset is not None:
            claims = preset[i] if isinstance(preset, list) and i < len(preset) else preset
            if isinstance(claims, dict):
                claims = claims.get("claims", [])
        rep = check_passage(p, retriever=retriever, k=args.k,
                            dry_run=args.dry_run, claims=claims)
        reports.append(rep)
        if not args.json:
            print("=" * 78)
            print(p.strip()[:300])
            print("=" * 78)
            render(rep, verbose=args.verbose)
            print()

    if args.json:
        print(json.dumps(reports, indent=2, ensure_ascii=False))
    if args.out:
        Path(args.out).write_text(json.dumps(reports, indent=2, ensure_ascii=False))
        print(f"wrote {args.out}", file=sys.stderr)


if __name__ == "__main__":
    sys.exit(main() or 0)
